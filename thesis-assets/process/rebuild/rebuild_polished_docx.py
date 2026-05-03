from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Any

from docx import Document


BASE = Path(__file__).resolve().parents[3]
ASSETS = BASE / "thesis-assets"
REBUILD = ASSETS / "process" / "rebuild"
REPORTS = ASSETS / "process" / "reports"

SOURCE_DOCX = ASSETS / "毕业论文正式版（草稿）.docx"
TARGET_DOCX = ASSETS / "current" / "毕业论文正式版（润色回填）.docx"
MARKDOWN = ASSETS / "current" / "thesis-polished.md"

WORD_STRUCTURE_JSON = REBUILD / "word-structure-map.json"
WORD_EXTRACTED_MD = REBUILD / "word-extracted-text.md"
MARKDOWN_STRUCTURE_JSON = REBUILD / "markdown-structure-map.json"
REBUILD_REPORT = REBUILD / "rebuild-report.md"
CHANGED_CSV = REBUILD / "changed-paragraphs.csv"
MISMATCH_REPORT = REBUILD / "mismatch-report.md"
CONTENT_CHECK = REPORTS / "polished-docx-content-check-2026-05-03.md"
FORMAT_PRECHECK = REPORTS / "format-precheck-after-rebuild-2026-05-03.md"

TOP_LEVEL = [
    "摘要",
    "Abstract",
    "第一章 绪论",
    "第二章 相关技术综述",
    "第三章 系统分析与设计",
    "第四章 系统实现",
    "第五章 系统测试",
    "第六章 总结与展望",
]

STOP_HEADINGS = {"参考文献", "致谢"}
FIGURE_REQUIRED = ["图3.1", "图3.2", "图3.3"]
TABLE_REQUIRED = ["表5.1", "表5.2", "表5.3", "表5.4"]


@dataclass
class ParagraphItem:
    chapter: str
    block: str
    paragraph_index: int | None
    style_name: str
    text: str
    kind: str
    is_heading: bool
    is_figure_caption: bool
    is_table_caption: bool
    level: int


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()


def norm(text: str) -> str:
    text = clean(text)
    text = re.sub(r"[`*_#]", "", text)
    text = re.sub(r"\s+", "", text)
    return text


def compact_number(text: str) -> str:
    return re.sub(r"([图表])\s+(\d+\.\d+)", r"\1\2", text)


def preview(text: str, limit: int = 90) -> str:
    text = clean(text)
    return text if len(text) <= limit else text[:limit] + "..."


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def para_text(paragraph) -> str:
    return clean("".join(run.text for run in paragraph.runs))


def set_paragraph_text(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def classify(text: str, style_name: str) -> tuple[str, bool, bool, bool, int]:
    is_heading = style_name.startswith("Heading")
    level = 0
    if style_name == "Heading 1":
        level = 1
    elif style_name == "Heading 2":
        level = 2
    elif text in TOP_LEVEL:
        is_heading = True
        level = 1
    elif re.match(r"^[1-6]\.\d+\s+", text):
        is_heading = True
        level = 2
    figure = bool(re.match(r"^图\s*\d+\.\d+\s+", text))
    table = bool(re.match(r"^表\s*\d+\.\d+\s+", text))
    kind = "heading" if is_heading else "caption" if figure or table else "body"
    return kind, is_heading, figure, table, level


def extract_word(docx_path: Path) -> tuple[dict[str, list[ParagraphItem]], list[dict[str, Any]]]:
    doc = Document(docx_path)
    sections: dict[str, list[ParagraphItem]] = {name: [] for name in TOP_LEVEL}
    raw: list[dict[str, Any]] = []
    current: str | None = None
    block = ""

    for idx, paragraph in enumerate(doc.paragraphs, 1):
        text = para_text(paragraph)
        style = paragraph.style.name if paragraph.style else ""
        if not text:
            continue

        if text in TOP_LEVEL:
            current = text
            block = text
        elif style.startswith("Heading") and (text in STOP_HEADINGS or text.startswith("附录")):
            current = None
            block = ""

        kind, is_heading, figure, table, level = classify(text, style)
        if current:
            if level in (1, 2):
                block = text
            item = ParagraphItem(
                chapter=current,
                block=block or current,
                paragraph_index=idx,
                style_name=style,
                text=text,
                kind=kind,
                is_heading=is_heading,
                is_figure_caption=figure,
                is_table_caption=table,
                level=level,
            )
            sections[current].append(item)
            raw.append(asdict(item))
    return sections, raw


def parse_markdown(path: Path) -> tuple[dict[str, list[ParagraphItem]], list[dict[str, Any]]]:
    sections: dict[str, list[ParagraphItem]] = {name: [] for name in TOP_LEVEL}
    raw: list[dict[str, Any]] = []
    current: str | None = None
    block = ""

    for line_no, line in enumerate(path.read_text(encoding="utf-8").splitlines(), 1):
        text = line.strip()
        if not text:
            continue
        if text.startswith("|") or re.match(r"^\|?\s*:?-{3,}", text):
            continue
        if text.startswith("```"):
            continue

        level = 0
        is_heading = False
        if text.startswith("#"):
            m = re.match(r"^(#+)\s+(.*)$", text)
            if not m:
                continue
            level = len(m.group(1))
            text = clean(m.group(2))
            is_heading = True
            if text in TOP_LEVEL:
                current = text
                block = text
            elif current and level == 2:
                block = text
        elif current:
            text = clean(text)

        if current:
            kind, _, figure, table, inferred_level = classify(text, "")
            if is_heading:
                kind = "heading"
                inferred_level = 1 if text in TOP_LEVEL else 2 if level == 2 else level
            item = ParagraphItem(
                chapter=current,
                block=block or current,
                paragraph_index=line_no,
                style_name="markdown",
                text=text,
                kind=kind,
                is_heading=is_heading,
                is_figure_caption=figure,
                is_table_caption=table,
                level=inferred_level,
            )
            sections[current].append(item)
            raw.append(asdict(item))
    return sections, raw


def grouped(items: list[ParagraphItem]) -> dict[str, list[ParagraphItem]]:
    groups: dict[str, list[ParagraphItem]] = {}
    for item in items:
        groups.setdefault(item.block, []).append(item)
    return groups


def write_structure_outputs(word_raw: list[dict[str, Any]], md_raw: list[dict[str, Any]]) -> None:
    WORD_STRUCTURE_JSON.write_text(json.dumps(word_raw, ensure_ascii=False, indent=2), encoding="utf-8")
    MARKDOWN_STRUCTURE_JSON.write_text(json.dumps(md_raw, ensure_ascii=False, indent=2), encoding="utf-8")
    lines = ["# Word 正文抽取", ""]
    for item in word_raw:
        marker = "#" if item["level"] == 1 else "##" if item["level"] == 2 else ""
        if marker:
            lines.append(f"{marker} {item['text']}")
        else:
            lines.append(item["text"])
        lines.append("")
    WORD_EXTRACTED_MD.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def section_matchable(word_items: list[ParagraphItem], md_items: list[ParagraphItem]) -> tuple[bool, str]:
    if len(word_items) != len(md_items):
        return False, f"paragraph_count word={len(word_items)} markdown={len(md_items)}"
    for w, m in zip(word_items, md_items):
        if w.kind != m.kind:
            return False, f"kind_mismatch word={w.kind} markdown={m.kind}"
        if w.kind == "heading" and w.level != m.level:
            return False, f"heading_level_mismatch word={w.level} markdown={m.level}"
        if w.is_figure_caption != m.is_figure_caption or w.is_table_caption != m.is_table_caption:
            return False, "caption_type_mismatch"
        if (w.is_figure_caption or w.is_table_caption) and compact_number(w.text).split(" ", 1)[0] != compact_number(m.text).split(" ", 1)[0]:
            return False, "caption_number_mismatch"
    return True, "ok"


def replace_docx_text(
    word_sections: dict[str, list[ParagraphItem]],
    md_sections: dict[str, list[ParagraphItem]],
) -> tuple[list[dict[str, Any]], list[dict[str, Any]], dict[str, int]]:
    doc = Document(SOURCE_DOCX)
    changed: list[dict[str, Any]] = []
    mismatches: list[dict[str, Any]] = []
    counts = {name: 0 for name in TOP_LEVEL}

    for chapter in TOP_LEVEL:
        w_groups = grouped(word_sections[chapter])
        m_groups = grouped(md_sections[chapter])
        for block_name, w_items in w_groups.items():
            if block_name not in m_groups:
                mismatches.append({
                    "chapter": chapter,
                    "issue_type": "missing_markdown_block",
                    "markdown_text_preview": "",
                    "word_text_preview": preview(w_items[0].text if w_items else block_name),
                    "suggested_action": "需人工处理：Markdown 中缺少对应小节。",
                })
                continue
            m_items = m_groups[block_name]
            ok, reason = section_matchable(w_items, m_items)
            if not ok:
                mismatches.append({
                    "chapter": chapter,
                    "issue_type": reason,
                    "markdown_text_preview": preview(" | ".join(i.text for i in m_items[:3])),
                    "word_text_preview": preview(" | ".join(i.text for i in w_items[:3])),
                    "suggested_action": "需人工处理：该小节段落结构不一致，已保留 Word 原内容。",
                })
                continue

            for w, m in zip(w_items, m_items):
                if w.paragraph_index is None:
                    continue
                if w.text == m.text:
                    continue
                paragraph = doc.paragraphs[w.paragraph_index - 1]
                set_paragraph_text(paragraph, m.text)
                changed.append({
                    "chapter": chapter,
                    "word_paragraph_index": w.paragraph_index,
                    "style_name": w.style_name,
                    "old_text_preview": preview(w.text),
                    "new_text_preview": preview(m.text),
                    "change_type": w.kind,
                })
                counts[chapter] += 1

        extra_md_blocks = sorted(set(m_groups) - set(w_groups))
        for block_name in extra_md_blocks:
            m_items = m_groups[block_name]
            mismatches.append({
                "chapter": chapter,
                "issue_type": "extra_markdown_block",
                "markdown_text_preview": preview(" | ".join(i.text for i in m_items[:3])),
                "word_text_preview": "",
                "suggested_action": "需人工处理：Markdown 中存在 Word 母版没有的小节或段落块。",
            })

    doc.save(TARGET_DOCX)
    return changed, mismatches, counts


def write_changed_csv(changed: list[dict[str, Any]]) -> None:
    with CHANGED_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(
            f,
            fieldnames=[
                "chapter",
                "word_paragraph_index",
                "style_name",
                "old_text_preview",
                "new_text_preview",
                "change_type",
            ],
        )
        writer.writeheader()
        writer.writerows(changed)


def write_mismatch_report(mismatches: list[dict[str, Any]]) -> None:
    lines = ["# mismatch-report", ""]
    if not mismatches:
        lines.append("- 无无法匹配段落。")
    else:
        lines.extend(["| chapter | issue_type | markdown_text_preview | word_text_preview | suggested_action |", "| --- | --- | --- | --- | --- |"])
        for item in mismatches:
            row = [
                item["chapter"],
                item["issue_type"],
                item["markdown_text_preview"].replace("|", "/"),
                item["word_text_preview"].replace("|", "/"),
                item["suggested_action"],
            ]
            lines.append("| " + " | ".join(row) + " |")
    MISMATCH_REPORT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def extract_after_sections(path: Path) -> dict[str, list[str]]:
    sections, _ = extract_word(path)
    return {k: [i.text for i in v] for k, v in sections.items()}


def markdown_sections_text(md_sections: dict[str, list[ParagraphItem]]) -> dict[str, list[str]]:
    return {k: [i.text for i in v] for k, v in md_sections.items()}


def content_check(md_sections: dict[str, list[ParagraphItem]], mismatches: list[dict[str, Any]]) -> dict[str, Any]:
    source_hash = sha256(SOURCE_DOCX)
    target_hash = sha256(TARGET_DOCX)
    target_same_as_source = source_hash == target_hash
    target_sections = extract_after_sections(TARGET_DOCX)
    md_text = markdown_sections_text(md_sections)
    chapter_results = {}
    all_major = True

    for chapter in TOP_LEVEL:
        target_join = norm("".join(target_sections.get(chapter, [])))
        md_items = [i for i in md_text.get(chapter, []) if i and not i.startswith("|")]
        md_join = norm("".join(md_items))
        if not md_join:
            contains_ratio = 0.0
        else:
            hits = 0
            total = 0
            for text in md_items:
                n = norm(text)
                if len(n) < 8:
                    continue
                total += 1
                if n in target_join:
                    hits += 1
            contains_ratio = hits / total if total else 1.0
        ok = contains_ratio >= 0.80
        if not ok:
            all_major = False
        chapter_results[chapter] = {
            "contains_ratio": round(contains_ratio, 3),
            "ok": ok,
        }

    target_all_text = "\n".join(t for vals in target_sections.values() for t in vals)
    citations = re.findall(r"\[(\d+(?:\s*[-,，]\s*\d+)*)\]", target_all_text)
    figures = set(re.findall(r"图\s*3\.\d+", target_all_text))
    tables = set(re.findall(r"表\s*5\.\d+", target_all_text))
    figure_ok = all(req in {compact_number(x) for x in figures} for req in FIGURE_REQUIRED)
    table_ok = all(req in {compact_number(x) for x in tables} for req in TABLE_REQUIRED)
    no_missing_chapters = all(target_sections.get(c) for c in TOP_LEVEL)
    success = (not target_same_as_source) and all_major and no_missing_chapters and figure_ok and table_ok and not mismatches

    return {
        "source_hash": source_hash,
        "target_hash": target_hash,
        "target_same_as_source": target_same_as_source,
        "chapter_results": chapter_results,
        "no_missing_chapters": no_missing_chapters,
        "citation_count": len(citations),
        "figure_ok": figure_ok,
        "table_ok": table_ok,
        "has_mismatches": bool(mismatches),
        "success": success,
    }


def write_rebuild_report(counts: dict[str, int], changed: list[dict[str, Any]], mismatches: list[dict[str, Any]]) -> None:
    lines = [
        "# DOCX 真实回填报告",
        "",
        f"- 输入 Word 母版：`{SOURCE_DOCX.relative_to(BASE).as_posix()}`",
        f"- 输入 Markdown 定稿：`{MARKDOWN.relative_to(BASE).as_posix()}`",
        f"- 输出 Word 工作稿：`{TARGET_DOCX.relative_to(BASE).as_posix()}`",
        "",
        "## 每章回填段落数",
        "",
        "| 章节 | 回填段落数 |",
        "| --- | ---: |",
    ]
    for chapter in TOP_LEVEL:
        lines.append(f"| {chapter} | {counts.get(chapter, 0)} |")
    lines.extend([
        "",
        "## 跳过区域",
        "",
        "- 封面、诚信责任书、目录、参考文献、致谢、附录。",
        "- Word 表格对象、图片对象、公式对象未展开替换。",
        "",
        "## 保留对象",
        "",
        "- 保留 Word 原有页眉页脚、页码、分节符和目录域。",
        "- 保留 Word 原有图片、表格和公式对象位置。",
        "",
        "## 核心结论",
        "",
        f"- 是否修改参考文献：否。",
        f"- 是否修改图表对象：否。",
        f"- 是否发现无法匹配段落：{'是' if mismatches else '否'}。",
        f"- 是否可以进入内容一致性校验：是。",
        f"- changed-paragraphs 记录数：{len(changed)}。",
    ])
    REBUILD_REPORT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_content_check_report(result: dict[str, Any]) -> None:
    lines = [
        "# 润色回填版 DOCX 内容一致性校验（2026-05-03）",
        "",
        f"- 原始主稿 SHA256：`{result['source_hash']}`",
        f"- 润色回填版 SHA256：`{result['target_hash']}`",
        f"- 润色回填版是否仍与原始主稿完全一致：{'是' if result['target_same_as_source'] else '否'}",
        "",
        "## 章节回填检查",
        "",
        "| 章节 | 与 Markdown 主要内容一致率 | 是否通过 |",
        "| --- | ---: | --- |",
    ]
    for chapter, item in result["chapter_results"].items():
        lines.append(f"| {chapter} | {item['contains_ratio']:.3f} | {'是' if item['ok'] else '否'} |")
    lines.extend([
        "",
        "## 编号检查",
        "",
        f"- 摘要、Abstract、第一章至第六章是否均存在：{'是' if result['no_missing_chapters'] else '否'}",
        f"- 文献引用编号数量：{result['citation_count']}",
        f"- 图表编号是否异常：{'否' if result['figure_ok'] and result['table_ok'] else '是'}",
        f"- 是否存在无法匹配段落：{'是' if result['has_mismatches'] else '否'}",
        "",
        "## 当前结论",
        "",
    ])
    if result["success"]:
        lines.append("- 可进入格式终审。")
    else:
        lines.append("- 需要人工处理 mismatch 后再次回填。")
    CONTENT_CHECK.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_format_precheck(result: dict[str, Any]) -> None:
    doc = Document(TARGET_DOCX)
    texts = [para_text(p) for p in doc.paragraphs if para_text(p)]
    joined = "\n".join(texts)
    headings = [t for t in texts if t in TOP_LEVEL or re.match(r"^[1-6]\.\d+\s+", t)]
    markdown_residue = []
    for marker in ["#", "**", "```"]:
        if marker in joined:
            markdown_residue.append(marker)
    if re.search(r"`[^`]+`", joined):
        markdown_residue.append("inline_backticks")
    empty_runs = sum(1 for p in doc.paragraphs if not para_text(p))
    placeholder_patterns = ["待补充", "TODO", "PLACEHOLDER", "这里填写"]
    placeholders = [p for p in placeholder_patterns if p in joined]
    compact_joined = compact_number(joined)

    lines = [
        "# 回填后格式预检查（2026-05-03）",
        "",
        f"- 是否存在摘要、Abstract、第一章至第六章：{'是' if all(c in texts for c in TOP_LEVEL) else '否'}",
        f"- 标题层级是否明显异常：{'否' if len(headings) >= 20 else '是'}",
        f"- 图题是否包含图3.1、图3.2、图3.3：{'是' if all(x in compact_joined for x in FIGURE_REQUIRED) else '否'}",
        f"- 表题是否包含表5.1 至 表5.4：{'是' if all(x in compact_joined for x in TABLE_REQUIRED) else '否'}",
        f"- 空段数量：{empty_runs}",
        f"- 是否存在明显 Markdown 残留符号：{'是：' + ', '.join(markdown_residue) if markdown_residue else '否'}",
        f"- 是否存在未替换占位文字：{'是：' + ', '.join(placeholders) if placeholders else '否'}",
        "",
        "说明：本检查不更新 Word 目录域，不导出 PDF，不做最终格式修复。",
    ]
    FORMAT_PRECHECK.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    REBUILD.mkdir(parents=True, exist_ok=True)
    REPORTS.mkdir(parents=True, exist_ok=True)
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)

    word_sections, word_raw = extract_word(SOURCE_DOCX)
    md_sections, md_raw = parse_markdown(MARKDOWN)
    write_structure_outputs(word_raw, md_raw)
    changed, mismatches, counts = replace_docx_text(word_sections, md_sections)
    write_changed_csv(changed)
    write_mismatch_report(mismatches)
    write_rebuild_report(counts, changed, mismatches)
    result = content_check(md_sections, mismatches)
    write_content_check_report(result)
    if result["success"]:
        write_format_precheck(result)
    return 0 if result["success"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
