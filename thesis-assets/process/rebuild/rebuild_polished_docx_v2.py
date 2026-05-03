from __future__ import annotations

import csv
import hashlib
import json
import re
import shutil
from copy import deepcopy
from dataclasses import asdict, dataclass
from pathlib import Path
from typing import Any

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


BASE = Path(__file__).resolve().parents[3]
ASSETS = BASE / "thesis-assets"
REBUILD = ASSETS / "process" / "rebuild"
REPORTS = ASSETS / "process" / "reports"

SOURCE_DOCX = ASSETS / "毕业论文正式版（草稿）.docx"
TARGET_DOCX = ASSETS / "current" / "毕业论文正式版（润色回填）.docx"
MARKDOWN = ASSETS / "current" / "thesis-polished.md"
FAILED_BACKUP = REBUILD / "毕业论文正式版（润色回填）-failed-partial-2026-05-03.docx"

SCRIPT = REBUILD / "rebuild_polished_docx_v2.py"
REBUILD_REPORT = REBUILD / "rebuild-report-v2.md"
CHANGED_CSV = REBUILD / "changed-paragraphs-v2.csv"
MISMATCH_REPORT = REBUILD / "mismatch-report-v2.md"
SECTION_SUMMARY_CSV = REBUILD / "section-rebuild-summary-v2.csv"
CONTENT_CHECK = REPORTS / "polished-docx-content-check-v2-2026-05-03.md"
FORMAT_PRECHECK = REPORTS / "format-precheck-after-rebuild-v2-2026-05-03.md"

TOP_LEVEL = [
    "摘要",
    "Abstract",
    "第一章 绪论",
    "第二章 相关技术综述",
    "第三章 系统分析与设计",
    "第四章 系统实现",
    "第五章 系统测试",
    "第六章 总结与展望",
]
STOP_HEADINGS = {"参考文献", "致谢"}
FOCUS_BLOCKS = [
    "2.2 Server-Sent Events 技术",
    "2.4 Spring Boot 3 异步编程模型",
    "2.5 Apache PDFBox 3.0",
    "4.3 简历解析模块实现",
    "5.3 功能测试",
    "6.1 总结",
    "6.2 不足与展望",
]
FIGURE_REQUIRED = ["图3.1", "图3.2", "图3.3"]
TABLE_REQUIRED = ["表5.1", "表5.2", "表5.3", "表5.4"]


@dataclass
class MdBlock:
    chapter: str
    title: str
    level: int
    paragraphs: list[str]


@dataclass
class WordBlock:
    chapter: str
    title: str
    level: int
    start: int
    end: int


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()


def strip_markdown(text: str) -> str:
    text = clean(text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = text.replace("**", "").replace("__", "")
    text = text.replace("*", "")
    return clean(text)


def norm(text: str) -> str:
    text = strip_markdown(text)
    text = re.sub(r"[`*_#]", "", text)
    return re.sub(r"\s+", "", text)


def compact_number(text: str) -> str:
    return re.sub(r"([图表])\s+(\d+\.\d+)", r"\1\2", text)


def preview(text: str, limit: int = 90) -> str:
    text = clean(text)
    return text if len(text) <= limit else text[:limit] + "..."


def sha256(path: Path) -> str:
    h = hashlib.sha256()
    with path.open("rb") as f:
        for chunk in iter(lambda: f.read(1024 * 1024), b""):
            h.update(chunk)
    return h.hexdigest()


def el_text(el) -> str:
    return clean("".join(t.text or "" for t in el.iter(qn("w:t"))))


def p_style_id(p) -> str:
    p_pr = p.find(qn("w:pPr"))
    if p_pr is None:
        return ""
    style = p_pr.find(qn("w:pStyle"))
    return style.get(qn("w:val")) if style is not None else ""


def is_p(el) -> bool:
    return el.tag == qn("w:p")


def is_tbl(el) -> bool:
    return el.tag == qn("w:tbl")


def has_tag(el, tag: str) -> bool:
    return any(x.tag == tag for x in el.iter())


def has_drawing(el) -> bool:
    return has_tag(el, qn("w:drawing")) or has_tag(el, qn("w:pict"))


def has_math(el) -> bool:
    return has_tag(el, qn("m:oMath")) or has_tag(el, qn("m:oMathPara"))


def has_field(el) -> bool:
    return has_tag(el, qn("w:fldChar")) or has_tag(el, qn("w:instrText"))


def heading_level(text: str, style_id: str) -> int:
    if text in TOP_LEVEL:
        return 1
    if text in STOP_HEADINGS or text.startswith("附录"):
        return 1
    if re.match(r"^[1-6]\.\d+\s+", text):
        return 2
    if style_id in {"Heading1", "1", "3"} and (text in TOP_LEVEL or re.match(r"^第[一二三四五六]章", text)):
        return 1
    if style_id in {"Heading2", "2", "4"} and re.match(r"^[1-6]\.\d+\s+", text):
        return 2
    return 0


def is_figure_caption(text: str) -> bool:
    return bool(re.match(r"^图\s*\d+\.\d+\s+", text))


def is_table_caption(text: str) -> bool:
    return bool(re.match(r"^表\s*\d+\.\d+\s+", text))


def is_safe_text_paragraph(p) -> bool:
    text = el_text(p)
    if not text:
        return False
    if has_drawing(p) or has_math(p) or has_field(p):
        return False
    if heading_level(text, p_style_id(p)):
        return False
    return True


def para_texts(docx_path: Path) -> list[str]:
    doc = Document(docx_path)
    return [clean("".join(run.text for run in p.runs)) for p in doc.paragraphs if clean("".join(run.text for run in p.runs))]


def set_paragraph_text(p, text: str) -> None:
    text = strip_markdown(text)
    runs = list(p.iter(qn("w:r")))
    if runs:
        first = True
        for r in runs:
            for child in list(r):
                if child.tag in {qn("w:t"), qn("w:tab"), qn("w:br")}:
                    r.remove(child)
            if first:
                t = OxmlElement("w:t")
                t.text = text
                r.append(t)
                first = False
    else:
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.text = text
        r.append(t)
        p.append(r)


def paragraph_kind(text: str) -> str:
    if is_figure_caption(text):
        return "figure_caption"
    if is_table_caption(text):
        return "table_caption"
    return "body"


def parse_markdown() -> tuple[dict[str, list[MdBlock]], list[dict[str, Any]]]:
    blocks: dict[str, list[MdBlock]] = {name: [] for name in TOP_LEVEL}
    raw: list[dict[str, Any]] = []
    current_chapter: str | None = None
    current_block: MdBlock | None = None

    def start_block(chapter: str, title: str, level: int) -> MdBlock:
        block = MdBlock(chapter=chapter, title=title, level=level, paragraphs=[])
        blocks[chapter].append(block)
        return block

    for line_no, line in enumerate(MARKDOWN.read_text(encoding="utf-8").splitlines(), 1):
        text = line.strip()
        if not text:
            continue
        if text.startswith("|") or re.match(r"^\|?\s*:?-{3,}", text):
            continue
        if text.startswith("```"):
            continue
        m = re.match(r"^(#+)\s+(.*)$", text)
        if m:
            level = len(m.group(1))
            title = strip_markdown(m.group(2))
            if level == 1 and title in TOP_LEVEL:
                current_chapter = title
                current_block = start_block(title, title, 1)
            elif level == 2 and current_chapter:
                current_block = start_block(current_chapter, title, 2)
            raw.append({"line": line_no, "kind": "heading", "level": level, "text": title, "chapter": current_chapter})
            continue
        if current_block:
            paragraph = strip_markdown(text)
            current_block.paragraphs.append(paragraph)
            raw.append({
                "line": line_no,
                "kind": paragraph_kind(paragraph),
                "level": 0,
                "text": paragraph,
                "chapter": current_block.chapter,
                "block": current_block.title,
            })
    return blocks, raw


def build_word_blocks(doc: Document) -> tuple[dict[str, list[WordBlock]], list[dict[str, Any]]]:
    body = doc.element.body
    children = list(body)
    headings: list[tuple[int, str, int]] = []
    raw: list[dict[str, Any]] = []
    current_chapter: str | None = None

    for idx, el in enumerate(children):
        if not is_p(el):
            continue
        text = el_text(el)
        if not text:
            continue
        level = heading_level(text, p_style_id(el))
        if level:
            if level == 1 and (text in STOP_HEADINGS or text.startswith("附录")):
                current_chapter = None
                headings.append((idx, text, level))
                raw.append({"body_index": idx, "text": text, "level": level, "chapter": current_chapter})
                continue
            if text in TOP_LEVEL:
                current_chapter = text
                headings.append((idx, text, level))
            elif level == 2 and current_chapter:
                headings.append((idx, text, level))
            raw.append({"body_index": idx, "text": text, "level": level, "chapter": current_chapter})

    blocks: dict[str, list[WordBlock]] = {name: [] for name in TOP_LEVEL}
    for i, (start, title, level) in enumerate(headings):
        chapter = title if title in TOP_LEVEL else None
        if level == 2:
            for prev_start, prev_title, prev_level in reversed(headings[:i]):
                if prev_level == 1 and prev_title in TOP_LEVEL:
                    chapter = prev_title
                    break
        if chapter not in blocks:
            continue
        end = len(children)
        for next_start, next_title, next_level in headings[i + 1:]:
            if next_level <= level:
                end = next_start
                break
        blocks[chapter].append(WordBlock(chapter=chapter, title=title, level=level, start=start, end=end))
    return blocks, raw


def safe_paragraphs_in_range(body, start: int, end: int) -> list[Any]:
    children = list(body)
    return [el for el in children[start + 1:end] if is_p(el) and is_safe_text_paragraph(el)]


def object_counts_in_range(body, start: int, end: int) -> dict[str, int]:
    children = list(body)
    p_nodes = [el for el in children[start:end] if is_p(el)]
    return {
        "tables": sum(1 for el in children[start:end] if is_tbl(el)),
        "pictures": sum(1 for el in p_nodes if has_drawing(el)),
        "formulas": sum(1 for el in p_nodes if has_math(el)),
    }


def insert_paragraph_before(body, anchor, template, text: str):
    new_p = deepcopy(template)
    set_paragraph_text(new_p, text)
    body.insert(body.index(anchor), new_p)
    return new_p


def append_paragraph_before_end(body, end_index: int, template, text: str):
    children = list(body)
    anchor = children[end_index] if end_index < len(children) else None
    new_p = deepcopy(template)
    set_paragraph_text(new_p, text)
    if anchor is None:
        body.append(new_p)
    else:
        body.insert(body.index(anchor), new_p)
    return new_p


def delete_paragraph(p) -> None:
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def style_name_for_docx_p(doc: Document, p) -> str:
    sid = p_style_id(p)
    if not sid:
        return ""
    try:
        return doc.styles[sid].name
    except Exception:
        return sid


def find_md_block(md_blocks: dict[str, list[MdBlock]], chapter: str, title: str) -> MdBlock | None:
    for block in md_blocks.get(chapter, []):
        if block.title == title:
            return block
    return None


def rebuild_docx(md_blocks: dict[str, list[MdBlock]]) -> tuple[list[dict[str, Any]], list[dict[str, Any]], list[dict[str, Any]], dict[str, int], dict[str, Any]]:
    shutil.copy2(SOURCE_DOCX, TARGET_DOCX)
    doc = Document(TARGET_DOCX)
    before_counts = count_objects(doc)
    changed: list[dict[str, Any]] = []
    mismatches: list[dict[str, Any]] = []
    summary: list[dict[str, Any]] = []
    chapter_counts = {name: 0 for name in TOP_LEVEL}

    word_blocks, _ = build_word_blocks(doc)
    for chapter in TOP_LEVEL:
        for wb in list(word_blocks.get(chapter, [])):
            md = find_md_block(md_blocks, chapter, wb.title)
            if md is None:
                mismatches.append({
                    "chapter": chapter,
                    "section": wb.title,
                    "issue_type": "missing_markdown_section",
                    "detail": "Markdown 中找不到同名章节或小节标题。",
                    "suggested_action": "人工确认标题是否被重命名。",
                })
                continue

            word_blocks, _ = build_word_blocks(doc)
            live_wb = find_word_block(word_blocks, chapter, wb.title)
            if live_wb is None:
                mismatches.append({
                    "chapter": chapter,
                    "section": wb.title,
                    "issue_type": "word_section_boundary_lost",
                    "detail": "插删后无法重新定位 Word 小节边界。",
                    "suggested_action": "人工检查小节标题样式和文本。",
                })
                continue

            body = doc.element.body
            children = list(body)
            heading = children[live_wb.start]
            old_heading = el_text(heading)
            if old_heading != md.title:
                set_paragraph_text(heading, md.title)
                changed.append(change_row(chapter, live_wb.title, live_wb.start, style_name_for_docx_p(doc, heading), old_heading, md.title, "heading"))
                chapter_counts[chapter] += 1

            if live_wb.level == 1 and live_wb.title not in {"摘要", "Abstract"}:
                continue

            safe_ps = safe_paragraphs_in_range(body, live_wb.start, live_wb.end)
            original_safe_count = len(safe_ps)
            md_count = len(md.paragraphs)
            counts = object_counts_in_range(body, live_wb.start, live_wb.end)
            inserted = 0
            deleted = 0
            completed = True

            if safe_ps:
                template = safe_ps[0]
            else:
                template = deepcopy(heading)
                set_paragraph_text(template, "")

            common = min(len(safe_ps), len(md.paragraphs))
            for i in range(common):
                old = el_text(safe_ps[i])
                new = md.paragraphs[i]
                if old != new:
                    set_paragraph_text(safe_ps[i], new)
                    changed.append(change_row(chapter, live_wb.title, live_wb.start, style_name_for_docx_p(doc, safe_ps[i]), old, new, paragraph_kind(new)))
                    chapter_counts[chapter] += 1

            if md_count > original_safe_count:
                word_blocks, _ = build_word_blocks(doc)
                live_wb = find_word_block(word_blocks, chapter, wb.title)
                if live_wb is None:
                    completed = False
                    mismatches.append({
                        "chapter": chapter,
                        "section": wb.title,
                        "issue_type": "insert_boundary_lost",
                        "detail": "插入前无法重新定位小节边界。",
                        "suggested_action": "人工检查小节标题。",
                    })
                else:
                    body = doc.element.body
                    for text in md.paragraphs[original_safe_count:]:
                        append_paragraph_before_end(body, live_wb.end, template, text)
                        changed.append(change_row(chapter, live_wb.title, live_wb.end, style_name_for_docx_p(doc, template), "", text, "insert_" + paragraph_kind(text)))
                        inserted += 1
                        chapter_counts[chapter] += 1

            if original_safe_count > md_count:
                for p in safe_ps[md_count:]:
                    old = el_text(p)
                    delete_paragraph(p)
                    changed.append(change_row(chapter, live_wb.title, live_wb.start, style_name_for_docx_p(doc, p), old, "", "delete_body"))
                    deleted += 1
                    chapter_counts[chapter] += 1

            summary.append({
                "chapter": chapter,
                "section": wb.title,
                "word_paragraph_count": original_safe_count,
                "markdown_paragraph_count": md_count,
                "inserted_paragraph_count": inserted,
                "deleted_paragraph_count": deleted,
                "preserved_object_count": counts["tables"] + counts["pictures"] + counts["formulas"],
                "tables": counts["tables"],
                "pictures": counts["pictures"],
                "formulas": counts["formulas"],
                "completed": "yes" if completed else "no",
                "focus_mismatch_section": "yes" if wb.title in FOCUS_BLOCKS else "no",
            })

    doc.save(TARGET_DOCX)
    after_counts = count_objects(Document(TARGET_DOCX))
    object_state = {
        "before": before_counts,
        "after": after_counts,
        "preserved": all(after_counts[k] >= before_counts[k] for k in before_counts),
    }
    return changed, mismatches, summary, chapter_counts, object_state


def find_word_block(word_blocks: dict[str, list[WordBlock]], chapter: str, title: str) -> WordBlock | None:
    for block in word_blocks.get(chapter, []):
        if block.title == title:
            return block
    return None


def count_objects(doc: Document) -> dict[str, int]:
    body = list(doc.element.body)
    p_nodes = [el for el in body if is_p(el)]
    return {
        "tables": len(doc.tables),
        "pictures": sum(1 for p in p_nodes if has_drawing(p)),
        "formulas": sum(1 for p in p_nodes if has_math(p)),
    }


def change_row(chapter: str, section: str, body_index: int, style_name: str, old: str, new: str, change_type: str) -> dict[str, Any]:
    return {
        "chapter": chapter,
        "section": section,
        "word_body_index": body_index,
        "style_name": style_name,
        "old_text_preview": preview(old),
        "new_text_preview": preview(new),
        "change_type": change_type,
    }


def write_changed_csv(changed: list[dict[str, Any]]) -> None:
    with CHANGED_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=[
            "chapter",
            "section",
            "word_body_index",
            "style_name",
            "old_text_preview",
            "new_text_preview",
            "change_type",
        ])
        writer.writeheader()
        writer.writerows(changed)


def write_summary_csv(summary: list[dict[str, Any]]) -> None:
    with SECTION_SUMMARY_CSV.open("w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=[
            "chapter",
            "section",
            "word_paragraph_count",
            "markdown_paragraph_count",
            "inserted_paragraph_count",
            "deleted_paragraph_count",
            "preserved_object_count",
            "tables",
            "pictures",
            "formulas",
            "completed",
            "focus_mismatch_section",
        ])
        writer.writeheader()
        writer.writerows(summary)


def write_mismatch_report(mismatches: list[dict[str, Any]]) -> None:
    lines = ["# mismatch-report-v2", ""]
    if not mismatches:
        lines.append("- 无真正无法处理的问题。")
    else:
        lines.extend([
            "| chapter | section | issue_type | detail | suggested_action |",
            "| --- | --- | --- | --- | --- |",
        ])
        for item in mismatches:
            lines.append("| " + " | ".join([
                item["chapter"],
                item["section"],
                item["issue_type"],
                item["detail"].replace("|", "/"),
                item["suggested_action"].replace("|", "/"),
            ]) + " |")
    MISMATCH_REPORT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def section_text_from_docx(path: Path) -> dict[str, dict[str, list[str]]]:
    doc = Document(path)
    blocks, _ = build_word_blocks(doc)
    result: dict[str, dict[str, list[str]]] = {name: {} for name in TOP_LEVEL}
    body = doc.element.body
    children = list(body)
    for chapter, word_blocks in blocks.items():
        for block in word_blocks:
            texts = []
            for el in children[block.start:block.end]:
                if is_p(el):
                    text = el_text(el)
                    if text:
                        texts.append(text)
            result[chapter][block.title] = texts
    return result


def md_text_map(md_blocks: dict[str, list[MdBlock]]) -> dict[str, dict[str, list[str]]]:
    return {
        chapter: {block.title: [block.title] + block.paragraphs for block in blocks}
        for chapter, blocks in md_blocks.items()
    }


def ratio_for_texts(md_items: list[str], target_join: str) -> tuple[float, list[str]]:
    hits = 0
    total = 0
    missing: list[str] = []
    for text in md_items:
        n = norm(text)
        if len(n) < 8:
            continue
        total += 1
        if n in target_join:
            hits += 1
        else:
            missing.append(preview(text))
    return (hits / total if total else 1.0), missing


def content_check(md_blocks: dict[str, list[MdBlock]], mismatches: list[dict[str, Any]], object_state: dict[str, Any]) -> dict[str, Any]:
    source_hash = sha256(SOURCE_DOCX)
    target_hash = sha256(TARGET_DOCX)
    target_same_as_source = source_hash == target_hash
    target_sections = section_text_from_docx(TARGET_DOCX)
    md_sections = md_text_map(md_blocks)
    chapter_results: dict[str, Any] = {}
    section_results: dict[str, Any] = {}

    for chapter in TOP_LEVEL:
        md_items = [x for block in md_sections.get(chapter, {}).values() for x in block]
        target_join = norm("".join(x for block in target_sections.get(chapter, {}).values() for x in block))
        ratio, missing = ratio_for_texts(md_items, target_join)
        chapter_results[chapter] = {"ratio": ratio, "ok": ratio >= 0.90, "missing": missing[:8]}

    for title in FOCUS_BLOCKS:
        chapter = next((c for c, blocks in md_sections.items() if title in blocks), "")
        md_items = md_sections.get(chapter, {}).get(title, [])
        target_items = target_sections.get(chapter, {}).get(title, [])
        ratio, missing = ratio_for_texts(md_items, norm("".join(target_items)))
        section_results[title] = {"chapter": chapter, "ratio": ratio, "ok": ratio >= 0.90, "missing": missing[:8]}

    target_text = "\n".join(para_texts(TARGET_DOCX))
    source_sections = section_text_from_docx(SOURCE_DOCX)
    target_full = section_text_from_docx(TARGET_DOCX)
    reference_unchanged = tail_region_text(SOURCE_DOCX) == tail_region_text(TARGET_DOCX)
    citations = re.findall(r"\[(\d+(?:\s*[-,，]\s*\d+)*)\]", target_text)
    compact = compact_number(target_text)
    figure_ok = all(x in compact for x in FIGURE_REQUIRED)
    table_ok = all(x in compact for x in TABLE_REQUIRED)
    residue = []
    for marker in ["#", "**", "```"]:
        if marker in target_text:
            residue.append(marker)
    if re.search(r"`[^`]+`", target_text):
        residue.append("inline_backticks")
    all_chapters = all(target_sections.get(chapter) for chapter in TOP_LEVEL)
    all_chapter_ok = all(item["ok"] for item in chapter_results.values())
    focus_ok = all(item["ok"] for item in section_results.values())
    success = (
        not target_same_as_source
        and all_chapters
        and all_chapter_ok
        and focus_ok
        and not mismatches
        and figure_ok
        and table_ok
        and not residue
        and reference_unchanged
        and object_state["preserved"]
    )
    return {
        "source_hash": source_hash,
        "target_hash": target_hash,
        "target_same_as_source": target_same_as_source,
        "chapter_results": chapter_results,
        "focus_results": section_results,
        "all_chapters": all_chapters,
        "citation_count": len(citations),
        "figure_ok": figure_ok,
        "table_ok": table_ok,
        "markdown_residue": residue,
        "reference_ack_appendix_unchanged": reference_unchanged,
        "object_state": object_state,
        "has_mismatches": bool(mismatches),
        "success": success,
        "source_section_count": sum(len(v) for v in source_sections.values()),
        "target_section_count": sum(len(v) for v in target_full.values()),
    }


def tail_region_text(path: Path) -> str:
    texts = para_texts(path)
    tail = []
    started = False
    for text in texts:
        if text == "参考文献":
            started = True
        if started:
            tail.append(text)
    return "\n".join(tail)


def write_content_check_report(result: dict[str, Any]) -> None:
    lines = [
        "# 润色回填版 DOCX 内容一致性校验 v2（2026-05-03）",
        "",
        f"- 原始主稿 SHA256：`{result['source_hash']}`",
        f"- 润色回填版 SHA256：`{result['target_hash']}`",
        f"- 润色回填版是否仍与原始主稿完全一致：{'是' if result['target_same_as_source'] else '否'}",
        f"- 摘要、Abstract、第一章至第六章是否均存在：{'是' if result['all_chapters'] else '否'}",
        "",
        "## 章节一致率",
        "",
        "| 章节 | 一致率 | 是否通过 | 主要差异 |",
        "| --- | ---: | --- | --- |",
    ]
    for chapter, item in result["chapter_results"].items():
        lines.append(f"| {chapter} | {item['ratio']:.3f} | {'是' if item['ok'] else '否'} | {'；'.join(item['missing']) if item['missing'] else ''} |")
    lines.extend([
        "",
        "## 上轮 mismatch 小节复核",
        "",
        "| 小节 | 一致率 | 是否解决 | 主要差异 |",
        "| --- | ---: | --- | --- |",
    ])
    for title, item in result["focus_results"].items():
        lines.append(f"| {title} | {item['ratio']:.3f} | {'是' if item['ok'] else '否'} | {'；'.join(item['missing']) if item['missing'] else ''} |")
    lines.extend([
        "",
        "## 编号与对象检查",
        "",
        f"- 文献引用编号数量：{result['citation_count']}",
        f"- 图题编号是否完整：{'是' if result['figure_ok'] else '否'}",
        f"- 表题编号是否完整：{'是' if result['table_ok'] else '否'}",
        f"- 是否存在 Markdown 残留符号：{'是：' + ', '.join(result['markdown_residue']) if result['markdown_residue'] else '否'}",
        f"- 参考文献、致谢、附录是否未被误改：{'是' if result['reference_ack_appendix_unchanged'] else '否'}",
        f"- 图片、表格、公式对象是否未减少：{'是' if result['object_state']['preserved'] else '否'}",
        f"- 对象计数：before={json.dumps(result['object_state']['before'], ensure_ascii=False)}, after={json.dumps(result['object_state']['after'], ensure_ascii=False)}",
        "",
        "## 当前结论",
        "",
        "- 可进入 Word 格式终审。" if result["success"] else "- 需要继续处理 v2 mismatch 或格式残留后再次回填。",
    ])
    CONTENT_CHECK.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_format_precheck(result: dict[str, Any]) -> None:
    texts = para_texts(TARGET_DOCX)
    joined = "\n".join(texts)
    compact = compact_number(joined)
    headings = [t for t in texts if t in TOP_LEVEL or re.match(r"^[1-6]\.\d+\s+", t)]
    duplicates = sorted({h for h in headings if headings.count(h) > 1})
    expected_sections = [block for block in FOCUS_BLOCKS]
    missing_sections = [s for s in expected_sections if s not in texts]
    residue = []
    for marker in ["#", "**", "```"]:
        if marker in joined:
            residue.append(marker)
    if re.search(r"`[^`]+`", joined):
        residue.append("inline_backticks")
    empty_count = sum(1 for p in Document(TARGET_DOCX).paragraphs if not clean("".join(r.text for r in p.runs)))
    lines = [
        "# 回填后格式预检查 v2（2026-05-03）",
        "",
        f"- 是否存在摘要、Abstract、第一章至第六章：{'是' if all(c in texts for c in TOP_LEVEL) else '否'}",
        f"- 标题层级是否存在明显异常：{'否' if len(headings) >= 20 else '是'}",
        f"- 图题编号是否存在图3.1、图3.2、图3.3：{'是' if all(x in compact for x in FIGURE_REQUIRED) else '否'}",
        f"- 表题编号是否存在表5.1 至 表5.4：{'是' if all(x in compact for x in TABLE_REQUIRED) else '否'}",
        f"- 明显空段数量：{empty_count}",
        f"- 是否存在 Markdown 残留符号：{'是：' + ', '.join(residue) if residue else '否'}",
        f"- 是否存在小节重复：{'是：' + '；'.join(duplicates) if duplicates else '否'}",
        f"- 是否存在重点小节缺失：{'是：' + '；'.join(missing_sections) if missing_sections else '否'}",
        "",
        "说明：本检查不更新 Word 目录域，不导出 PDF，不做最终格式修复。",
    ]
    FORMAT_PRECHECK.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_rebuild_report(chapter_counts: dict[str, int], summary: list[dict[str, Any]], mismatches: list[dict[str, Any]], object_state: dict[str, Any], content_result: dict[str, Any]) -> None:
    focus_status = {row["section"]: row for row in summary if row["section"] in FOCUS_BLOCKS}
    inserted_total = sum(int(row["inserted_paragraph_count"]) for row in summary)
    deleted_total = sum(int(row["deleted_paragraph_count"]) for row in summary)
    lines = [
        "# DOCX 真实回填 v2 报告",
        "",
        f"- 输入 Word 母版：`{SOURCE_DOCX.relative_to(BASE).as_posix()}`",
        f"- 输入 Markdown 定稿：`{MARKDOWN.relative_to(BASE).as_posix()}`",
        f"- 输出 Word 工作稿：`{TARGET_DOCX.relative_to(BASE).as_posix()}`",
        f"- 本轮是否从原始母版重新生成：是",
        f"- 是否覆盖上一轮失败产物：是，目标文件已从原始母版重新生成；上一轮失败产物本地备份为 `{FAILED_BACKUP.relative_to(BASE).as_posix()}`，默认不提交。",
        "",
        "## 每章回填段落数",
        "",
        "| 章节 | 回填变更数 |",
        "| --- | ---: |",
    ]
    for chapter in TOP_LEVEL:
        lines.append(f"| {chapter} | {chapter_counts.get(chapter, 0)} |")
    lines.extend([
        "",
        "## 上轮 mismatch 小节解决状态",
        "",
        "| 小节 | Word 原段落数 | Markdown 段落数 | 插入段落数 | 删除段落数 | 保留对象数 | 是否完成回填 |",
        "| --- | ---: | ---: | ---: | ---: | ---: | --- |",
    ])
    for title in FOCUS_BLOCKS:
        row = focus_status.get(title)
        if row:
            lines.append(f"| {title} | {row['word_paragraph_count']} | {row['markdown_paragraph_count']} | {row['inserted_paragraph_count']} | {row['deleted_paragraph_count']} | {row['preserved_object_count']} | {row['completed']} |")
        else:
            lines.append(f"| {title} |  |  |  |  |  | no |")
    lines.extend([
        "",
        "## 核心结论",
        "",
        f"- 是否删除普通文本段落：{'是' if deleted_total else '否'}，共 {deleted_total} 段。",
        f"- 是否插入普通文本段落：{'是' if inserted_total else '否'}，共 {inserted_total} 段。",
        f"- 是否保留表格、图片、公式对象：{'是' if object_state['preserved'] else '否'}。",
        f"- 对象计数：before={json.dumps(object_state['before'], ensure_ascii=False)}, after={json.dumps(object_state['after'], ensure_ascii=False)}。",
        "- 是否修改参考文献：否。",
        "- 是否修改附录：否。",
        f"- v2 是否仍有真正无法处理的问题：{'是' if mismatches else '否'}。",
        f"- 内容一致性是否通过：{'是' if content_result['success'] else '否'}。",
    ])
    REBUILD_REPORT.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    REBUILD.mkdir(parents=True, exist_ok=True)
    REPORTS.mkdir(parents=True, exist_ok=True)
    if TARGET_DOCX.exists():
        shutil.copy2(TARGET_DOCX, FAILED_BACKUP)

    md_blocks, md_raw = parse_markdown()
    (REBUILD / "markdown-structure-map-v2.json").write_text(json.dumps(md_raw, ensure_ascii=False, indent=2), encoding="utf-8")
    changed, mismatches, summary, chapter_counts, object_state = rebuild_docx(md_blocks)
    write_changed_csv(changed)
    write_summary_csv(summary)
    write_mismatch_report(mismatches)
    content_result = content_check(md_blocks, mismatches, object_state)
    write_content_check_report(content_result)
    if content_result["success"]:
        write_format_precheck(content_result)
    write_rebuild_report(chapter_counts, summary, mismatches, object_state, content_result)
    return 0 if content_result["success"] else 2


if __name__ == "__main__":
    raise SystemExit(main())
